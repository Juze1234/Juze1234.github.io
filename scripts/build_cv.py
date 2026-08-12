from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "public" / "cv"
OUTPUT_PATH = OUTPUT_DIR / "Sergey_Senchenko_CV.docx"

EDUCATION_INSTITUTION = "Tallinn Polytechnic School (Tallinna Polütehnikum)"
EDUCATION_PROGRAM = "IT Specialist"
EDUCATION_YEARS = ""  # set to e.g. "2018-2021" to show years

FONT = "Calibri"
INK = RGBColor(28, 31, 32)
MUTED = RGBColor(94, 101, 101)
ACCENT = RGBColor(157, 28, 28)
LIGHT = RGBColor(216, 220, 219)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_border_bottom(paragraph, color="9D1C1C", size=10, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink(paragraph, text, url, color="6B1D1D"):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), "19")
    r_pr.extend([r_fonts, color_el, size_el])
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    existing_abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    existing_num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract_ids, default=0) + 1
    num_id = max(existing_num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level_type)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "252")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract_num.append(lvl)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.insert(0, num_pr)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    h1 = document.styles["Heading 1"]
    h1.font.name = FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h1.font.size = Pt(11.5)
    h1.font.bold = True
    h1.font.color.rgb = ACCENT
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(3.5)
    h1.paragraph_format.keep_with_next = True

    h2 = document.styles["Heading 2"]
    h2.font.name = FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2.font.size = Pt(10.5)
    h2.font.bold = True
    h2.font.color.rgb = INK
    h2.paragraph_format.space_before = Pt(4)
    h2.paragraph_format.space_after = Pt(1)
    h2.paragraph_format.keep_with_next = True

    if "CV Metadata" not in [style.name for style in document.styles]:
        meta = document.styles.add_style("CV Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = document.styles["CV Metadata"]
    meta.font.name = FONT
    meta._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    meta.font.size = Pt(9.5)
    meta.font.color.rgb = MUTED
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.line_spacing = 1.0


def add_section_heading(document, text):
    p = document.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_run_font(run, size=11.5, color=ACCENT, bold=True)
    return p


def add_bullet(document, num_id, text):
    p = document.add_paragraph()
    apply_bullet(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=9.15, color=INK)
    return p


def add_skill_line(document, label, value):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.04
    label_run = p.add_run(label + ": ")
    set_run_font(label_run, size=9.15, color=INK, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=9.15, color=INK)


def build_cv():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    properties = document.core_properties
    properties.title = "Sergey Senchenko - Technical Designer CV"
    properties.subject = "Technical Designer, Gameplay Scripter, and Level Designer"
    properties.author = "Sergey Senchenko"
    properties.keywords = "Technical Design, Gameplay Scripting, Enforce Script, DayZ, Level Design, QA"

    # Header pattern: compact memo masthead, adapted for an ATS-friendly CV.
    name = document.add_paragraph()
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(1)
    run = name.add_run("SERGEY SENCHENKO")
    set_run_font(run, size=24, color=INK, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Technical Designer | Gameplay Scripter | Level Designer")
    set_run_font(run, size=11.5, color=ACCENT, bold=True)

    contact = document.add_paragraph(style="CV Metadata")
    contact.add_run("Tallinn, Estonia | Open to remote opportunities | ")
    add_hyperlink(contact, "SergeyJuze@gmail.com", "mailto:SergeyJuze@gmail.com")
    contact.add_run("\nPortfolio: ")
    add_hyperlink(contact, "juze1234.github.io", "https://juze1234.github.io/")
    contact.add_run(" | LinkedIn: ")
    add_hyperlink(contact, "linkedin.com/in/sergey-senchenko-951aa837b", "https://www.linkedin.com/in/sergey-senchenko-951aa837b/")
    contact.add_run(" | GitHub: ")
    add_hyperlink(contact, "github.com/Juze1234", "https://github.com/Juze1234")
    contact.add_run(" | Project videos: ")
    add_hyperlink(contact, "youtube.com/@MetroWARrpOFF", "https://www.youtube.com/@MetroWARrpOFF")
    set_paragraph_border_bottom(contact, color="9D1C1C", size=8, space=7)

    add_section_heading(document, "Professional Summary")
    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(3)
    summary.paragraph_format.line_spacing = 1.08
    run = summary.add_run(
        "Technical designer and gameplay scripter with seven years of development experience, the last two "
        "spent leading a multiplayer DayZ project end to end. Works across Enforce Script, server-authoritative "
        "gameplay systems, data-driven player customization, persistent identity, progression and economy, "
        "level design, asset integration, and hands-on QA. Comfortable owning a feature from concept through "
        "implementation, testing, and player-facing iteration."
    )
    set_run_font(run, size=9.4, color=INK)

    add_section_heading(document, "Experience")
    role = document.add_paragraph(style="Heading 2")
    role.paragraph_format.space_after = Pt(0)
    run = role.add_run("Lead Developer / Technical Designer - Metro W.A.R. RP")
    set_run_font(run, size=10.5, color=INK, bold=True)
    meta = document.add_paragraph(style="CV Metadata")
    meta.paragraph_format.space_after = Pt(3)
    run = meta.add_run("Independent DayZ project | 2024-Present | Two-person team")
    set_run_font(run, size=9.1, color=MUTED, italic=True)

    num_id = create_bullet_numbering(document)
    bullets = [
        "Lead development across gameplay scripting, technical design, model integration, game design, multiplayer economy, and QA; collaborate with a second developer who focuses on world and map production.",
        "Designed and implemented a data-driven character loadout system with real-time preview, persistent profile rules, and server-side eligibility validation.",
        "Built synchronized interactive world systems: power networks, switches, lighting circuits, control panels, powered doors, and the player feedback around them.",
        "Developed persistent player identity, in-game economy, access control, private storage, environmental hazards, and progression systems.",
        "Built a rideable rail vehicle for multiplayer, covering physics tuning, multi-passenger support, custom track layouts, and automatic alignment to valid track nodes.",
        "Contributed level design across underground and surface environments, shaping onboarding, exploration routes, economy touchpoints, atmosphere, and encounter flow.",
        "Own QA alongside development: reproduce issues, trace client-server failures, validate configuration edge cases, and iterate through multiplayer playtests and player feedback.",
    ]
    for bullet in bullets:
        add_bullet(document, num_id, bullet)

    impact = document.add_paragraph()
    impact.paragraph_format.space_before = Pt(2)
    impact.paragraph_format.space_after = Pt(2)
    impact.paragraph_format.line_spacing = 1.03
    label = impact.add_run("PROJECT IMPACT  ")
    set_run_font(label, size=8.8, color=ACCENT, bold=True)
    value = impact.add_run(
        "2 years in development | 900+ community members | Full 80-player server during the latest major test | Next closed beta in preparation"
    )
    set_run_font(value, size=8.8, color=INK, bold=True)

    add_section_heading(document, "Skills")
    add_skill_line(document, "Programming", "Gameplay scripting (Enforce Script); client-server architecture and RPC; data persistence; config-driven systems; debugging")
    add_skill_line(document, "Game development", "Technical design; gameplay systems; multiplayer design; game economy; level design; onboarding; hands-on QA")
    add_skill_line(document, "Tools", "DayZ Tools; Object Builder; Terrain Builder; Blender; Substance Painter; Adobe Photoshop")

    add_section_heading(document, "Education")
    education = document.add_paragraph()
    education.paragraph_format.space_after = Pt(0)
    run = education.add_run(
        f"{EDUCATION_INSTITUTION} - {EDUCATION_PROGRAM}"
        + (f" | {EDUCATION_YEARS}" if EDUCATION_YEARS else "")
    )
    set_run_font(run, size=9.2, color=INK)

    add_section_heading(document, "Languages")
    languages = document.add_paragraph()
    languages.paragraph_format.space_after = Pt(0)
    run = languages.add_run("Russian - Native | English - Professional working proficiency | Estonian - Basic")
    set_run_font(run, size=9.2, color=INK)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("Portfolio source code remains private; a technical walkthrough is available during recruitment.")
    set_run_font(run, size=7.5, color=MUTED)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_cv()
